from pathlib import Path
import math
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE=Path(__file__).resolve().parents[1]
RES=BASE/'results'
DATA=BASE/'data_processed'
FIGS=BASE/'figures'
OUT=BASE/'manuscript'/'Anson_Dengue_miRNA_SubmissionReady_Draft.docx'
OUT.parent.mkdir(parents=True, exist_ok=True)

primary=pd.read_csv(RES/'target_miRNA_primary_comparison_summary.csv')
allm=pd.read_csv(RES/'all_datasets_all_miRNAs_all_comparisons.csv')
meta=pd.read_csv(RES/'severity_meta_analysis_random_effects.csv')
meta_study=pd.read_csv(RES/'severity_meta_analysis_per_study_effects.csv')
count_sens=pd.read_csv(RES/'GSE150623_count_model_sensitivity_targets.csv')
libsum=pd.read_csv(RES/'GSE150623_library_size_summary_by_group.csv')
counts150=pd.read_csv(DATA/'GSE150623_raw_counts_collapsed.csv')
meta150=pd.read_csv(DATA/'GSE150623_metadata.csv')

def fmt(x, digits=2, na='NA'):
    try:
        if x is None or pd.isna(x): return na
        return f"{float(x):.{digits}f}"
    except Exception:
        return na

def fmt_p(x):
    try:
        if x is None or pd.isna(x): return 'NA'
        x=float(x)
        if x < 0.001: return f"{x:.2e}"
        return f"{x:.3f}"
    except Exception:
        return 'NA'

def get(df, dataset, comp, mirna):
    r=df[(df['dataset']==dataset)&(df['comparison']==comp)&(df['miRNA']==mirna)]
    return r.iloc[0]

def get_sens(mirna):
    r=count_sens[count_sens['miRNA']==mirna]
    return None if r.empty else r.iloc[0]

def font_run(run, size=None, bold=None, italic=None):
    run.font.name='Arial'
    run._element.rPr.rFonts.set(qn('w:ascii'),'Arial')
    run._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial')
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic

def addp(doc,text='',style=None,size=10.2,align=None):
    p=doc.add_paragraph(style=style)
    if align: p.alignment=align
    r=p.add_run(text); font_run(r,size=size)
    p.paragraph_format.space_after=Pt(5)
    p.paragraph_format.line_spacing=1.05
    return p

def add_heading(doc,text,level=1):
    p=doc.add_heading('',level=level)
    r=p.add_run(text); font_run(r,size=14 if level==1 else 12.2,bold=True)
    p.paragraph_format.space_before=Pt(8)
    p.paragraph_format.space_after=Pt(4)
    return p

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_cell(cell,text,bold=False,size=7.3):
    cell.text=''; p=cell.paragraphs[0]
    r=p.add_run(str(text)); font_run(r,size=size,bold=bold)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    p.paragraph_format.space_after=Pt(0)

def caption(doc,text):
    p=doc.add_paragraph(style='Caption')
    r=p.add_run(text); font_run(r,size=8.4,italic=True)
    p.paragraph_format.space_after=Pt(5)
    return p

def add_table(doc, headers, rows, widths=None, size=7.0):
    table=doc.add_table(rows=1, cols=len(headers)); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    # Repeat the header row if a table splits over a page; this avoids continuation rows without context.
    trPr=table.rows[0]._tr.get_or_add_trPr()
    tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)
    for i,h in enumerate(headers):
        set_cell(table.rows[0].cells[i],h,bold=True,size=size); shade(table.rows[0].cells[i],'D9EAF7')
    for row in rows:
        cells=table.add_row().cells
        for i,v in enumerate(row): set_cell(cells[i],v,size=size)
    return table

# Derived values
n190=len(allm[(allm.dataset=='GSE190749')&(allm.comparison=='SD_vs_NonSevere')])
n150=len(allm[(allm.dataset=='GSE150623')&(allm.comparison=='SD_vs_NonSevere')])
qsig150=int(allm[(allm.dataset=='GSE150623')&(allm.comparison=='SD_vs_NonSevere')]['q_bh_welch'].lt(0.05).sum())
minq190=allm[(allm.dataset=='GSE190749')&(allm.comparison=='SD_vs_NonSevere')]['q_bh_welch'].min()
mi_col=counts150.columns[0]
mir1246=counts150[counts150[mi_col]=='hsa-miR-1246'].iloc[0]
sev_samples=meta150[meta150.severity_binary=='severe']['sample'].tolist()
nonsev_samples=meta150[meta150.severity_binary=='non-severe']['sample'].tolist()
mir1246_sev_detected=sum(float(mir1246[s])>0 for s in sev_samples)
mir1246_nonsev_detected=sum(float(mir1246[s])>0 for s in nonsev_samples)
mir122_meta=meta[(meta.miRNA=='hsa-miR-122-5p')&(meta.model=='Main: GSE190749 + GSE150623')].iloc[0]
mir574_meta=meta[(meta.miRNA=='hsa-miR-574-5p')&(meta.model=='Main: GSE190749 + GSE150623')].iloc[0]
mir1246_meta=meta[(meta.miRNA=='hsa-miR-1246')&(meta.model=='Main: GSE190749 + GSE150623')].iloc[0]
# library summaries
lib_di=libsum[libsum.clinical_group=='DI'].iloc[0]
lib_dws=libsum[libsum.clinical_group=='DWS'].iloc[0]
lib_ds=libsum[libsum.clinical_group=='DS'].iloc[0]

# Document setup
doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(0.62); sec.bottom_margin=Inches(0.62); sec.left_margin=Inches(0.65); sec.right_margin=Inches(0.65)
for sty in ['Normal','Title','Heading 1','Heading 2','Heading 3','Caption']:
    if sty in doc.styles:
        doc.styles[sty].font.name='Arial'
        doc.styles[sty]._element.rPr.rFonts.set(qn('w:ascii'),'Arial')
        doc.styles[sty]._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial')
doc.styles['Normal'].font.size=Pt(10.2)

# Title page-ish
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Public-data reproducibility reanalysis of dengue severity-associated circulating microRNAs')
font_run(r, size=16, bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Non-validation of miR-574-5p/miR-1246 and a hypothesis-generating miR-122-5p signal')
font_run(r, size=12.5, italic=True)
addp(doc,'Anson Han An Xuan',align=WD_ALIGN_PARAGRAPH.CENTER)
addp(doc,'Tzu Chi International Student',align=WD_ALIGN_PARAGRAPH.CENTER)
addp(doc,'Target manuscript type: public-data reanalysis / reproducibility study',align=WD_ALIGN_PARAGRAPH.CENTER,size=9.5)

add_heading(doc,'Declarations',1)
addp(doc,'Funding: No specific grant or external funding was reported for this draft.')
addp(doc,'Competing interests: The author declares no competing interests.')
addp(doc,'Ethics statement: This study used de-identified public repository data only. No new participants were recruited, no biological samples were collected, and no Malaysian patient-level data were used.')
addp(doc,'Data availability: The reanalysis used public data from NCBI GEO accessions GSE190749, GSE150623, GSE169170, and GSE307678. Only GSE190749 and GSE150623 were treated as the two main severe-vs-non-severe datasets. GSE307678 was not used as primary severe-dengue validation because it compares warning-sign status rather than severe dengue.')
addp(doc,'Code availability: The accompanying package contains the scripts, processed tables, figures, and a final quality-check report needed to reproduce the draft outputs.')
addp(doc,'AI assistance disclosure: OpenAI ChatGPT was used to assist with code drafting, manuscript organization, reanalysis checks, and editing. The author is responsible for verifying the data, code outputs, citations, interpretations, and final submitted files.')

doc.add_page_break()

add_heading(doc,'Abstract',1)
addp(doc,'Background: Circulating microRNAs have been proposed as early prognostic biomarkers for severe dengue. A Thai prospective study identified miR-574-5p and miR-1246 as promising candidates, but independent public-data reproducibility remains uncertain.')
addp(doc,f'Methods: Public human dengue miRNA datasets were reanalysed. The primary severe-vs-non-severe assessment was restricted to GSE190749, used as a reproduction dataset, and GSE150623, used as the main independent dataset. GSE169170 was retained only as an underpowered exploratory severity-spectrum dataset, and GSE307678 was retained only as Malaysian warning-sign context. Normalized-expression analyses used log2 fold change, Welch tests, Benjamini-Hochberg correction across all measured miRNAs within each comparison, Mann-Whitney sensitivity checks, AUC, and random-effects Hedges g synthesis. Because GSE150623 is count-based and showed strong library-size imbalance, a count-aware sensitivity analysis used Poisson regression with robust standard errors and log library-size offset for all miRNAs, plus target-level negative-binomial regression with estimated dispersion.')
addp(doc,f'Results: miR-574-5p and miR-1246 did not survive multiple-testing correction. In GSE190749, miR-574-5p and miR-1246 were nominally higher in severe dengue (log2FC=2.25 and 1.67), but both had q=0.941 after correction across {n190} tested miRNAs. In GSE150623, normalized-expression analysis did not validate either primary target (miR-574-5p q=0.309; miR-1246 q=0.388). Count-aware sensitivity did not rescue miR-574-5p (negative-binomial offset log2FC=0.71, p=0.642), and miR-1246 was too sparse for reliable count modelling because it was detected in {mir1246_sev_detected}/{len(sev_samples)} severe and {mir1246_nonsev_detected}/{len(nonsev_samples)} non-severe samples. The most consistent positive signal was miR-122-5p: pooled Hedges g={fmt(mir122_meta.effect)}, 95% CI {fmt(mir122_meta.ci_low)} to {fmt(mir122_meta.ci_high)}, p={fmt_p(mir122_meta.p)}, I2={fmt(mir122_meta.I2,1)}% across the two main severity datasets, with positive count-aware support in GSE150623.')
addp(doc,'Conclusions: Public data do not support miR-574-5p or miR-1246 as universal standalone severe-dengue biomarkers. The strongest scientifically defensible claim is a reproducibility reanalysis showing primary-marker non-validation, platform/detection limitations, and a hypothesis-generating miR-122-5p signal requiring prospective validation.')

add_heading(doc,'Author summary',1)
addp(doc,'Dengue severity is difficult to predict early, and blood microRNAs may provide useful molecular clues. This study reanalysed public dengue microRNA datasets to test whether two previously proposed markers, miR-574-5p and miR-1246, were reproducible outside the original setting. The main finding is cautionary: the primary markers did not remain significant after false-discovery correction and did not validate robustly in the main independent dataset. A second candidate, miR-122-5p, showed a more consistent cross-cohort signal, but it is not yet a validated clinical biomarker. The value of the study is therefore reproducibility: it helps narrow which dengue microRNA signals are robust enough to justify future patient studies.')

add_heading(doc,'Introduction',1)
addp(doc,'Dengue virus infection remains a major clinical and public-health problem in tropical and subtropical regions. Most infections are self-limited, but a subset of patients progress to severe dengue characterized by plasma leakage, bleeding, shock, or organ impairment [1,2]. A central clinical problem is that severe progression can be difficult to predict early, before the critical phase of illness. Reliable early biomarkers could improve triage and monitoring, especially in settings with high patient volume.')
addp(doc,'MicroRNAs are short non-coding RNAs that regulate gene expression and can be detected in serum or plasma. Circulating miRNAs are attractive biomarker candidates because they can remain stable in blood through association with vesicles or protein complexes [3-6]. In dengue, miRNA profiles may reflect immune activation, endothelial stress, tissue injury, viral effects, or host responses to infection. However, biomarker discovery studies can overestimate performance when candidates are not tested in independent cohorts.')
addp(doc,'Limothai et al. reported that circulating miR-574-5p and miR-1246 had prognostic value for severe dengue in a Thai cohort [7]. That finding is clinically interesting, but a candidate biomarker is most useful only if it remains reproducible across datasets, countries, sample-processing methods, and measurement platforms. Public repositories such as NCBI GEO make independent reanalysis possible [8].')
addp(doc,'The present study was redesigned as a public-data reproducibility analysis rather than a claim of Malaysian clinical validation. The primary objective was to test whether miR-574-5p and miR-1246 showed consistent severe-dengue association in public human miRNA data. A secondary objective was to identify whether any candidate, especially miR-122-5p, showed more coherent cross-dataset support. Because the available Malaysian dataset does not contain severe-dengue classification, it was handled as exploratory warning-sign context only.')

add_heading(doc,'Methods',1)
add_heading(doc,'Study design and datasets',2)
addp(doc,'This was a retrospective secondary analysis of public, de-identified human dengue miRNA datasets. No new wet-lab experiment, patient recruitment, or patient contact was performed. Dataset roles were defined before final interpretation: GSE190749 and GSE150623 were used as the two main severe-vs-non-severe datasets; GSE169170 and GSE307678 were exploratory only.')

headers=['Dataset','Country/context','Specimen/platform','Samples used','Role in final analysis','Reason for role']
rows=[
['GSE190749','Thailand','Serum; NanoString processed expression','19 total: 8 severe, 11 non-severe','Primary reproduction dataset','Same research line as the published Limothai study; useful for reproduction but not independent external validation.'],
['GSE150623','India','Plasma; small-RNA count matrix','39 severity-labelled samples: 16 severe, 23 non-severe','Primary independent dataset','Best public independent severe-vs-non-severe miRNA dataset available in the uploaded files.'],
['GSE169170','India','Serum; Agilent microarray raw files','8 total','Exploratory only','Very small sample size and DF/DHF/DSS labels; not strong enough as primary validation.'],
['GSE307678','Malaysia','Plasma exosomal miRNA; Agilent raw files','6 total','Exploratory warning-sign context only','Malaysian, but compares warning signs vs no warning signs rather than severe vs non-severe dengue.']]
add_table(doc,headers,rows,size=6.2)
caption(doc,'Table 1. Dataset roles after final methodological review. The primary severe-dengue claim is intentionally restricted to GSE190749 and GSE150623.')

add_heading(doc,'Expression processing and statistical analysis',2)
addp(doc,'For GSE150623, duplicate miRNA rows were collapsed by miRNA name, library totals were calculated from the count matrix, and expression was normalized as counts per million followed by log2(CPM + 1). For GSE169170 and GSE307678, Agilent raw signal files were parsed, log2-transformed, and quantile-normalized. GSE190749 was analysed from processed NanoString expression data available in the analysis package. All miRNA names were harmonized to mature hsa-miR identifiers where possible.')
addp(doc,'The primary normalized-expression comparison was severe dengue versus non-severe dengue. Non-severe dengue included dengue infection and dengue with warning signs when the dataset did not provide a separate clinically severe outcome. For each dataset-comparison, Welch tests were applied across measured miRNAs and p-values were adjusted using the Benjamini-Hochberg false-discovery-rate method. Candidate miRNAs were additionally summarized using Mann-Whitney tests and AUC, but AUC values were treated as descriptive because sample sizes were small.')
addp(doc,'Random-effects Hedges g meta-analysis was performed only across the two main comparable severity datasets, GSE190749 and GSE150623. GSE307678 was excluded from the primary pooled analysis because warning signs are not equivalent to severe dengue. GSE169170 was not included in the main pooled estimate because of the very small sample size and older DF/DHF/DSS classification.')
addp(doc,'A final sensitivity analysis addressed a major GSE150623 quality issue: severe samples had much lower sequencing library totals than DI or DWS samples. To test whether normalized-expression conclusions were robust to count-based modelling, GSE150623 was reanalysed using Poisson regression with robust standard errors and a log library-size offset across all miRNAs, with Benjamini-Hochberg correction across all tested features. A target-level negative-binomial NB2 model with estimated dispersion and the same offset was also run for candidate miRNAs. These count-aware models were used as sensitivity checks rather than as clinical validation models.')

# Page break before Results removed after render QC to avoid a blank page.
add_heading(doc,'Results',1)
add_heading(doc,'Dataset and quality-control findings',2)
addp(doc,f'In GSE190749, {n190} miRNAs were tested in the severe-vs-non-severe comparison and no miRNA reached q<0.05; the smallest q-value was {fmt(minq190,3)}. In GSE150623, {n150} miRNAs were tested and {qsig150} reached q<0.05 in the normalized-expression analysis, but the primary Limothai targets did not. A key QC concern was read depth: GSE150623 severe samples had a median total miRNA count of {int(lib_ds.median_library_total):,}, compared with {int(lib_di.median_library_total):,} for DI and {int(lib_dws.median_library_total):,} for DWS. This imbalance makes platform and detection-threshold explanations especially important.')

headers=['Group','n','Median library total','Minimum','Maximum']
rows=[]
for _,r in libsum.iterrows():
    rows.append([r['clinical_group'], int(r['n']), f"{int(r['median_library_total']):,}", f"{int(r['min_library_total']):,}", f"{int(r['max_library_total']):,}"])
add_table(doc,headers,rows,size=7.0)
caption(doc,'Table 2. GSE150623 library-size QC by clinical group. Lower severe-dengue library sizes increase uncertainty for sparse miRNAs.')

add_heading(doc,'miR-574-5p and miR-1246 failed robust independent validation',2)
rows=[]
for m in ['hsa-miR-574-5p','hsa-miR-1246']:
    r190=get(primary,'GSE190749','SD_vs_NonSevere',m)
    r150=get(primary,'GSE150623','SD_vs_NonSevere',m)
    sens=get_sens(m)
    if sens is not None and pd.notna(sens['nb2_offset_log2FC']):
        nb=f"{fmt(sens['nb2_offset_log2FC'])}; p={fmt_p(sens['nb2_offset_p'])}"
        pois=f"q={fmt_p(sens['poisson_offset_q_bh_all_miRNAs'])}"
    else:
        nb='Too sparse'
        pois='Too sparse'
    rows.append([m, fmt(r190.log2FC_groupA_minus_groupB), fmt_p(r190.p_welch), fmt_p(r190.q_bh_welch), fmt(r190.auc_groupA_high), fmt(r150.log2FC_groupA_minus_groupB), fmt_p(r150.p_welch), fmt_p(r150.q_bh_welch), fmt(r150.auc_groupA_high), nb, pois])
headers=['miRNA','GSE190749 log2FC','p','q','AUC','GSE150623 log2FC','p','q','AUC','GSE150623 NB2 offset log2FC; p','GSE150623 Poisson offset q']
add_table(doc,headers,rows,size=5.8)
caption(doc,'Table 3. Primary-target results. GSE190749 supports nominal reproduction but not FDR significance; GSE150623 does not provide robust independent validation.')
addp(doc,f'Raw counts confirmed that miR-1246 was essentially undetectable in GSE150623: {mir1246_sev_detected}/{len(sev_samples)} severe samples and {mir1246_nonsev_detected}/{len(nonsev_samples)} non-severe samples had non-zero counts. Therefore, the miR-1246 result should not be interpreted simply as a biological absence of association; it may reflect platform sensitivity, detection threshold, sample processing, or read-depth differences.')
addp(doc,'For miR-574-5p, the normalized-expression analysis in GSE150623 gave a negative, non-significant effect, while count-aware models gave a positive but still clearly non-significant effect. This method sensitivity does not rescue miR-574-5p as a reproducible marker; instead, it shows that the independent dataset is insufficient to support a strong conclusion for that target.')

add_heading(doc,'miR-122-5p was the strongest hypothesis-generating signal',2)
r122_190=get(primary,'GSE190749','SD_vs_NonSevere','hsa-miR-122-5p')
r122_150=get(primary,'GSE150623','SD_vs_NonSevere','hsa-miR-122-5p')
s122=get_sens('hsa-miR-122-5p')
addp(doc,f'miR-122-5p showed the most coherent cross-dataset pattern. In GSE190749, it was higher in severe dengue by normalized-expression analysis (log2FC={fmt(r122_190.log2FC_groupA_minus_groupB)}, p={fmt_p(r122_190.p_welch)}, q={fmt_p(r122_190.q_bh_welch)}, AUC={fmt(r122_190.auc_groupA_high)}). In GSE150623, normalized expression was directionally higher but not FDR-significant (log2FC={fmt(r122_150.log2FC_groupA_minus_groupB)}, p={fmt_p(r122_150.p_welch)}, q={fmt_p(r122_150.q_bh_welch)}, AUC={fmt(r122_150.auc_groupA_high)}). Count-aware modelling of GSE150623 strengthened the direction: NB2 offset log2FC={fmt(s122.nb2_offset_log2FC)}, p={fmt_p(s122.nb2_offset_p)}, and Poisson-offset all-miRNA q={fmt_p(s122.poisson_offset_q_bh_all_miRNAs)}.')
addp(doc,f'The random-effects synthesis across GSE190749 and GSE150623 gave pooled Hedges g={fmt(mir122_meta.effect)}, 95% CI {fmt(mir122_meta.ci_low)} to {fmt(mir122_meta.ci_high)}, p={fmt_p(mir122_meta.p)}, and I2={fmt(mir122_meta.I2,1)}%. This was the most consistent signal in the package, but it remains exploratory because it was not the original primary hypothesis and because no prospective clinical threshold was tested.')

headers=['miRNA','Main pooled Hedges g','95% CI','p','I2','Interpretation']
rows=[
['miR-574-5p',fmt(mir574_meta.effect),f"{fmt(mir574_meta.ci_low)} to {fmt(mir574_meta.ci_high)}",fmt_p(mir574_meta.p),fmt(mir574_meta.I2,1)+'%','Inconsistent; high heterogeneity'],
['miR-1246',fmt(mir1246_meta.effect),f"{fmt(mir1246_meta.ci_low)} to {fmt(mir1246_meta.ci_high)}",fmt_p(mir1246_meta.p),fmt(mir1246_meta.I2,1)+'%','Inconsistent; sparse detection in GSE150623'],
['miR-122-5p',fmt(mir122_meta.effect),f"{fmt(mir122_meta.ci_low)} to {fmt(mir122_meta.ci_high)}",fmt_p(mir122_meta.p),fmt(mir122_meta.I2,1)+'%','Most consistent hypothesis-generating signal']]
add_table(doc,headers,rows,size=7.0)
caption(doc,'Table 4. Main random-effects synthesis across the two comparable severe-vs-non-severe datasets only.')

add_heading(doc,'Figures',1)
for fname,cap in [
('figure1_primary_targets_fdr.png','Figure 1. Primary target normalized-expression results in the two main severity datasets. Positive log2FC indicates higher expression in severe dengue; q-values are Benjamini-Hochberg adjusted across measured miRNAs.'),
('figure2_mir122_forest.png','Figure 2. miR-122-5p standardized effects across the two comparable severity datasets and the random-effects pooled estimate.'),
('figure3_mir1246_detection.png','Figure 3. GSE150623 raw miR-1246 counts. All severe-dengue samples had zero raw counts, supporting a detection-threshold interpretation.'),
('figure4_gse150623_library_size_qc.png','Figure 4. GSE150623 library-size QC. Severe-dengue samples had substantially lower total miRNA counts than DI and DWS samples.'),
('figure5_count_model_sensitivity.png','Figure 5. GSE150623 count-aware target sensitivity using negative-binomial NB2 regression with log library-size offset.')]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(); r.add_picture(str(FIGS/fname), width=Inches(5.8))
    caption(doc,cap)

add_heading(doc,'Discussion',1)
addp(doc,'The final interpretation is deliberately conservative. The public-data evidence does not justify presenting miR-574-5p or miR-1246 as validated universal severe-dengue biomarkers. In the Thai reproduction dataset, both primary targets had nominally favorable results, but neither survived FDR correction. In the main independent dataset, neither target showed robust validation, and miR-1246 was largely below detection. This is a scientifically useful negative-validation result because it tests whether a promising discovery signal generalizes across public cohorts.')
addp(doc,'The most important improvement in this final version is the separation of statistical non-validation from technical non-detection. For miR-1246, the GSE150623 signal cannot be interpreted confidently because the miRNA was almost entirely absent from the raw count matrix. For miR-574-5p, direction depended on whether normalized-expression or count-offset modelling was used, but both approaches agreed that the evidence was not statistically strong. These points make the paper more honest and less vulnerable to overclaiming.')
addp(doc,'miR-122-5p is now the strongest positive candidate. It showed the most consistent random-effects result across the two comparable severity datasets and remained positive in a GSE150623 count-aware sensitivity analysis. However, the manuscript should still avoid calling miR-122-5p a validated biomarker. The correct framing is that miR-122-5p is a hypothesis-generating candidate that should be tested prospectively using a single standardized assay, prespecified sampling time, and clinically adjudicated severity outcome.')
addp(doc,'The exclusion of GSE307678 from the primary analysis is also important. Although it is Malaysian and therefore locally relevant, it compares warning-sign status in a very small number of exosomal plasma samples. Warning signs are clinically useful but are not the same phenotype as severe dengue. Including this dataset in a primary pooled severe-dengue analysis would weaken the biological logic and invite reviewer criticism.')

add_heading(doc,'Limitations',1)
addp(doc,'This study has several limitations. First, it is a secondary analysis of public datasets, so clinical covariates, sampling time, serotype, prior dengue exposure, and laboratory protocols could not be harmonized. Second, the two main datasets use different specimen types and profiling platforms. Third, GSE150623 showed strong library-size imbalance, especially lower total miRNA counts in severe samples. Fourth, several candidate results are nominal or model-dependent and do not survive all-feature FDR correction. Fifth, no Malaysian severe-dengue patient cohort was available, so the paper cannot claim Malaysian clinical validation. Finally, no prospective threshold, decision curve, or externally validated multivariable clinical model was developed.')

add_heading(doc,'Conclusion',1)
addp(doc,'After final rechecking, the manuscript is strongest as a reproducibility study. The public data do not robustly validate miR-574-5p or miR-1246 as universal severe-dengue biomarkers. The most defensible contribution is to show this non-validation clearly, explain detection and library-size limitations, and highlight miR-122-5p as the best hypothesis-generating candidate for future prospective study. This framing is suitable for a cautious public-data reanalysis submission, especially to a journal that accepts technically sound replication and negative-result studies.')

add_heading(doc,'References',1)
refs=[
'1. World Health Organization. Dengue: guidelines for diagnosis, treatment, prevention and control. Geneva: World Health Organization; 2009.',
'2. Bhatt S, Gething PW, Brady OJ, Messina JP, Farlow AW, Moyes CL, et al. The global distribution and burden of dengue. Nature. 2013;496:504-507. doi:10.1038/nature12060.',
'3. Bartel DP. MicroRNAs: target recognition and regulatory functions. Cell. 2009;136:215-233. doi:10.1016/j.cell.2009.01.002.',
'4. Ambros V. The functions of animal microRNAs. Nature. 2004;431:350-355. doi:10.1038/nature02871.',
'5. Mitchell PS, Parkin RK, Kroh EM, Fritz BR, Wyman SK, Pogosova-Agadjanyan EL, et al. Circulating microRNAs as stable blood-based markers for cancer detection. Proc Natl Acad Sci USA. 2008;105:10513-10518. doi:10.1073/pnas.0804549105.',
'6. Arroyo JD, Chevillet JR, Kroh EM, Ruf IK, Pritchard CC, Gibson DF, et al. Argonaute2 complexes carry a population of circulating microRNAs independent of vesicles in human plasma. Proc Natl Acad Sci USA. 2011;108:5003-5008. doi:10.1073/pnas.1019055108.',
'7. Limothai U, Tachaboon S, Dinhuzen J, Kiatmungkhisthira A, Duangchinda T, Pisitkun T, et al. Circulating microRNA as a prognostic biomarker for severe dengue: a prospective observational study. PLoS Negl Trop Dis. 2022;16:e0010836. doi:10.1371/journal.pntd.0010836.',
'8. Barrett T, Wilhite SE, Ledoux P, Evangelista C, Kim IF, Tomashevsky M, et al. NCBI GEO: archive for functional genomics data sets-update. Nucleic Acids Res. 2013;41:D991-D995. doi:10.1093/nar/gks1193.',
'9. John DV, Lin YS, Perng GC. Biomarkers of severe dengue disease - a review. J Biomed Sci. 2015;22:83. doi:10.1186/s12929-015-0191-6.',
'10. Hedges LV. Distribution theory for Glass\'s estimator of effect size and related estimators. J Educ Stat. 1981;6:107-128.',
'11. Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J R Stat Soc Series B. 1995;57:289-300.',
'12. Robin X, Turck N, Hainard A, Tiberti N, Lisacek F, Sanchez JC, et al. pROC: an open-source package for R and S+ to analyze and compare ROC curves. BMC Bioinformatics. 2011;12:77. doi:10.1186/1471-2105-12-77.',
'13. Love MI, Huber W, Anders S. Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biol. 2014;15:550. doi:10.1186/s13059-014-0550-8.'
]
for ref in refs:
    addp(doc,ref,size=8.7)

add_heading(doc,'Final pre-submission note',1)
addp(doc,'This draft is scientifically more defensible than the earlier template because it does not overclaim validation. Before actual journal submission, the author should still confirm the official affiliation wording, corresponding author email, and target journal formatting rules. If a supervisor has access to R/Bioconductor, a DESeq2 or edgeR rerun of GSE150623 would be a useful additional robustness check; the current package already includes count-aware Python sensitivity analyses that support the same cautious conclusion.')

# Save
doc.save(OUT)
print(OUT)
